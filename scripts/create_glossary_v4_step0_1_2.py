#!/usr/bin/env python3
"""
Create glossary.json for book1_quants/module1 (Rate and Return)
Following GLOSSARY_INSTRUCTION_v4.md STRICT ALGORITHM

ШАГ 0: Read ENTIRE PDF
ШАГ 1: Find ALL LOS patterns
ШАГ 2: Define LOS boundaries
ШАГ 3: Extract terms for EACH LOS
ШАГ 4: Add calculator for terms with formulas
ШАГ 5: Final validation
"""

import json
import re
from docx import Document
from pathlib import Path

# Paths
notes_docx = "/home/user/CFA-LVL-I-TRAINER/frontend/data/v2/book1_quants/module1/sources/notes.docx"
output_path = "/home/user/CFA-LVL-I-TRAINER/frontend/data/v2/book1_quants/module1/glossary.json"

print("="*80)
print("CREATING GLOSSARY FOR MODULE 1: RATE AND RETURN")
print("Following GLOSSARY_INSTRUCTION_v4.md")
print("="*80)

# ШАГ 0: READ ENTIRE DOCUMENT
print("\n📖 ШАГ 0: Reading ENTIRE document from beginning to end...")

doc = Document(notes_docx)
full_text = "\n".join([para.text for para in doc.paragraphs])
total_chars = len(full_text)
total_paras = len(doc.paragraphs)

print(f"   ✓ Total paragraphs: {total_paras}")
print(f"   ✓ Total characters: {total_chars:,}")
print(f"   ✓ Document fully loaded")

# ШАГ 1: FIND ALL LOS PATTERNS
print("\n🔍 ШАГ 1: Finding ALL LOS patterns using regex...")

# Pattern: "LOS \d+[a-z]:"
los_pattern = r'LOS\s+(\d+[a-z]):'
los_matches = re.findall(los_pattern, full_text, re.IGNORECASE)

# Remove duplicates and sort
los_found = sorted(list(set(los_matches)))
los_list = [f"LOS_{los}" for los in los_found]

print(f"   Found LOS patterns: {los_found}")
print(f"   Total LOS count: {len(los_list)}")
print(f"   los_list: {los_list}")

# STOP-CHECK #1
if len(los_list) < 3:
    print("\n⛔ STOP-CHECK #1 FAILED!")
    print(f"   Found only {len(los_list)} LOS (minimum required: 3)")
    print("   ERROR: Scanning failed - returning to STEP 0")
    exit(1)
else:
    print(f"   ✓ STOP-CHECK #1 PASSED: {len(los_list)} LOS found (≥3)")

# ШАГ 2: DEFINE LOS BOUNDARIES
print("\n📍 ШАГ 2: Defining boundaries for each LOS...")

los_boundaries = []
for i, los in enumerate(los_found):
    start_pattern = f"LOS {los}:"
    start_pos = full_text.find(start_pattern)

    # End position is the start of next LOS or end of document
    if i < len(los_found) - 1:
        next_los = los_found[i + 1]
        end_pattern = f"LOS {next_los}:"
        end_pos = full_text.find(end_pattern)
    else:
        end_pos = len(full_text)

    los_text = full_text[start_pos:end_pos]
    los_boundaries.append({
        "los_id": f"LOS_{los}",
        "los": los,
        "start": start_pos,
        "end": end_pos,
        "text": los_text,
        "length": len(los_text)
    })

    print(f"   LOS_{los}: chars {start_pos:,} to {end_pos:,} (length: {len(los_text):,})")

print(f"   ✓ All {len(los_boundaries)} LOS boundaries defined")

print("\n")
print("="*80)
print("DOCUMENT STRUCTURE ANALYSIS")
print("="*80)
for boundary in los_boundaries:
    preview = boundary['text'][:200].replace('\n', ' ')
    print(f"\n{boundary['los_id']}:")
    print(f"   Length: {boundary['length']:,} chars")
    print(f"   Preview: {preview}...")

print("\n\n⏸️  CHECKPOINT: Document structure analyzed")
print(f"✓ {len(los_list)} LOS found: {', '.join(los_list)}")
print("\nReady for ШАГ 3: Term extraction")
print("\nPress Enter to continue or Ctrl+C to stop...")
