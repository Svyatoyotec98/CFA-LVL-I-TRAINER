#!/usr/bin/env python3
"""
Read first 30 paragraphs from notes.docx to evaluate conversion quality
"""

from docx import Document
from pathlib import Path

docx_path = "/home/user/CFA-LVL-I-TRAINER/frontend/data/v2/book1_quants/module1/sources/notes.docx"

print(f"📖 Reading: {Path(docx_path).name}")
print(f"=" * 80)
print()

doc = Document(docx_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Showing first 30 paragraphs:\n")
print("=" * 80)

for i, para in enumerate(doc.paragraphs[:30], 1):
    text = para.text.strip()
    if text:  # Only show non-empty paragraphs
        print(f"\n[Para {i}]")
        print(text)
        print("-" * 80)

print("\n" + "=" * 80)
print(f"✅ Displayed first 30 paragraphs from notes.docx")
